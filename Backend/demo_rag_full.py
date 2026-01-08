from rag_engine import RAGEngine
import os
import shutil
import sys
import fitz # PyMuPDF
from docx import Document

# Force UTF-8
sys.stdout.reconfigure(encoding='utf-8')

# 1. Setup
kb_dir = "demo_final_kb"
if os.path.exists(kb_dir):
    shutil.rmtree(kb_dir) # Clean start

def create_files():
    # PDF
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), "Project Alpha Confidential.\nThe launch date is set for December 25th.\nCode: RED-EAGLE.\nPage 1.")
    doc.save("demo_alpha.pdf")
    
    # DOCX
    d = Document()
    d.add_paragraph("HR Policy (DOCX).")
    d.add_paragraph("Remote work is allowed on Fridays.")
    d.save("demo_hr.docx")
    
    return [os.path.abspath("demo_alpha.pdf"), os.path.abspath("demo_hr.docx")]

print(f"--- RAG System Demo (Storage: {kb_dir}) ---")

# 2. Initialize
rag = RAGEngine(use_openai_embeddings=False, storage_dir=kb_dir)

# 3. Ingest
files = create_files()
print(f"Ingesting {len(files)} documents...")
rag.load_documents(files)

# 4. Search 1 (PDF Content)
q1 = "What is the code for Project Alpha?"
print(f"\n[Q1]: '{q1}'")
results1 = rag.retrieve_relevant_chunks(q1)

# 5. Search 2 (DOCX Content)
q2 = "Can I work from home?"
print(f"\n[Q2]: '{q2}'")
results2 = rag.retrieve_relevant_chunks(q2)

# Cleanup
if os.path.exists(kb_dir):
    shutil.rmtree(kb_dir)
# Removing files
try:
    os.remove("demo_alpha.pdf")
    os.remove("demo_hr.docx")
except:
    pass
