import os
import sys

# Ensure UTF-8 output
sys.stdout.reconfigure(encoding='utf-8')

try:
    from docx import Document
except ImportError:
    print("python-docx not installed!")
    exit(1)

import fitz
from rag_engine import RAGEngine

def create_dummy_files():
    # 1. Create PDF
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), "Standard Operating Procedure (PDF).\nAlways wear safety gear.\nPage 1.")
    doc.save("demo_doc.pdf")
    
    # 2. Create DOCX
    d = Document()
    d.add_paragraph("Employee Handbook (DOCX).")
    d.add_paragraph("Policy 1: Work hours are 9 to 5.")
    d.save("demo_doc.docx")
    
    # 3. Create TXT
    with open("demo_doc.txt", "w", encoding="utf-8") as f:
        f.write("Meeting Notes (TXT).\nDiscussed project timeline.")

    print("Dummy files created.")

def test():
    create_dummy_files()
    
    print("\n--- Initializing Engine ---")
    rag = RAGEngine(use_openai_embeddings=False, storage_dir="demo_kb")
    
    print("\n--- Ingesting Files ---")
    files = [
        os.path.abspath("demo_doc.pdf"),
        os.path.abspath("demo_doc.docx"),
        os.path.abspath("demo_doc.txt")
    ]
    rag.load_documents(files)
    
    print("\n--- searching 'policy' ---")
    results = rag.retrieve_relevant_chunks("What are the work hours?")
    
    print("\n--- searching 'safety' ---")
    results2 = rag.retrieve_relevant_chunks("What gear should I wear?")

    # Check Persistence
    if os.path.exists("test_kb/knowledge_base.json"):
        print("\n[SUCCESS] knowledge_base.json found.")
    else:
        print("\n[FAIL] knowledge_base.json MISSING.")

if __name__ == "__main__":
    test()
